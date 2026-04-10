from __future__ import annotations

import io
from pathlib import Path

import fitz
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from loguru import logger
from pdf2docx import Converter as Pdf2DocxConverter

from app.services.ocr_service import extract_page_elements, extract_text_from_pdf, is_text_pdf
from app.services.validator import validate_pdf_file


class ConversionError(RuntimeError):
    """Exception raised when any stage of the PDF-to-DOCX conversion process fails."""


def _docx_contains_text(docx_path: Path) -> bool:
    """
    Check if the generated DOCX file contains any readable text blocks.
    
    Args:
        docx_path: Standard pathlib.Path wrapper indicating the DOCX output file.
        
    Returns:
        bool: True if paragraph text exists within the document space.
    """
    document = Document(docx_path)
    return any(paragraph.text.strip() for paragraph in document.paragraphs)


def _convert_pdf2docx(pdf_path: Path, output_path: Path) -> None:
    """
    Wraps the standard pdf2docx utility converter processing.
    
    Args:
        pdf_path: The local filesystem path of the source PDF.
        output_path: The local filesystem path representing the desired `.docx` destination.
    """
    with Pdf2DocxConverter(str(pdf_path)) as converter:
        converter.convert(str(output_path), start=0, end=None)


def _create_docx_from_ocr(pdf_path: Path, output_path: Path, lang: str) -> None:
    """
    Provides a last resort pure-OCR fallback when hybrid algorithms and typical 
    converters fail formatting evaluations.
    
    Args:
        pdf_path: Source input target PDF filepath.
        output_path: Storage output filepath representation.
        lang: Optical Character Recognition language constraint string type.
    """
    extracted_text: str = extract_text_from_pdf(pdf_path, lang=lang)
    if not extracted_text.strip():
        raise ConversionError("Secondary pure OCR extraction methodology yielded zero valid string blocks.")

    document = Document()
    for paragraph in extracted_text.split('\n\n'):
        paragraph = paragraph.strip()
        if paragraph:
            document.add_paragraph(paragraph)

    document.save(str(output_path))


def _clean_font_name(pdf_font_name: str) -> str:
    """
    Normalize font tags from PyMuPDF dict trees to simple root system names.
    e.g. converting `BAAAAA+SourceSansPro-Bold` directly to `SourceSansPro`.
    
    Args:
        pdf_font_name: The internally structured PDF font string array wrapper.
        
    Returns:
        str: Expected system font rendering name.
    """
    if not pdf_font_name:
        return "Arial"

    if "+" in pdf_font_name:
        pdf_font_name = pdf_font_name.split("+", 1)[1]

    suffixes: list[str] = [
        "-BoldItalic", "-Bold", "-Italic", "MT", "PS",
        "Bold", "Italic", "-Oblique", "Oblique"
    ]
    
    changed: bool = True
    while changed:
        changed = False
        for suffix in suffixes:
            if pdf_font_name.endswith(suffix):
                pdf_font_name = pdf_font_name[:-len(suffix)]
                changed = True
                break

    pdf_font_name = pdf_font_name.rstrip("-")
    
    if not pdf_font_name:
        return "Arial"

    return pdf_font_name


def _convert_hybrid(pdf_path: Path, output_path: Path, ocr_lang: str) -> None:
    """
    Premium conversion engine prioritizing accurate inline positioning of stylized Text, 
    preserving precise text elements (Font string, size constraint points, boldness layout) 
    and precisely calculated aspect-ratios of inline visual imagery using python-docx.
    
    Args:
        pdf_path: Original base PDF source node.
        output_path: Target directory constraint destination reference path.
        ocr_lang: Fallback execution framework linguistic instruction.
    """
    document = Document()
    
    with fitz.open(str(pdf_path)) as pdf_doc:
        total_pages: int = len(pdf_doc)
        if total_pages == 0:
            raise ConversionError("Evaluated physical PDF file presents 0 embedded pages.")
            
        for page_num in range(total_pages):
            page: fitz.Page = pdf_doc[page_num]
            logger.debug(f"Parsing advanced spatial dictionary blocks against internal tree for index {page_num + 1}/{total_pages}...")
            
            elements = extract_page_elements(page, ocr_lang=ocr_lang)
            text_count: int = sum(1 for e in elements if e[1] == "text")
            image_count: int = sum(1 for e in elements if e[1] == "image")
            
            logger.info(f"Page [{page_num + 1}] element topology loaded - Text spans discovered: {text_count}, Inline images discovered: {image_count}")
            
            for _, el_type, content in elements:
                if el_type == "text":
                    paragraph = document.add_paragraph()
                    for span in content:
                        run = paragraph.add_run(span["text"])
                        
                        # Font System Identification Name Extraction Mapping
                        font_name: str = _clean_font_name(span["font"])
                        if (span.get("flags", 0) & 8) or "Courier" in font_name or "Mono" in font_name:
                            font_name = "Courier New"
                        run.font.name = font_name
                        
                        # Apply Size Restraints Constraints Correctly
                        size: float = span["size"]
                        if size <= 0 or size is None:
                            size = 11.0
                        run.font.size = Pt(size)
                        
                        # Determine Stylistic Visual Weight Context (Bold)
                        if span["bold"] or "Bold" in span["font"]:
                            run.bold = True
                            
                        # Extract Angle Styling Restraints Matrix (Italic)
                        if span["italic"] or "Italic" in span["font"] or "Oblique" in span["font"]:
                            run.italic = True
                            
                        # Evaluate Accurate Base Text Visual Chroma Range Vector Mask Mapping Output Stream Rendering Value Colors
                        r, g, b = span["color_r"], span["color_g"], span["color_b"] # RGB Vector array extraction elements from dictionary array pointer
                        if not (r == 0 and g == 0 and b == 0):  # Pure standard zero output is dropped gracefully
                            run.font.color.rgb = RGBColor(r, g, b)
                            
                elif el_type == "image":
                    try:
                        image_bytes: bytes = content["bytes"]
                        img_width_pt: float = content["width_pt"]
                        page_width_pt: float = content["page_width_pt"]
                        
                        # Compute aspect calculation relative sizing to fit 6.5 docx standard layout pages constraint metrics accurately
                        ratio: float = img_width_pt / float(page_width_pt) if page_width_pt > 0 else 1.0
                        calc_wd: float = 6.5 * ratio
                        calc_wd = max(1.0, min(6.5, calc_wd))
                        
                        image_stream = io.BytesIO(image_bytes)
                        document.add_picture(image_stream, width=Inches(calc_wd))
                    except Exception as e:
                        logger.warning(f"Engine dynamically circumvented picture format layout insertion bypass stream structure constraint bounds mapping parsing error exception on layout view block page map {page_num + 1}: {e}")
            
            if page_num < total_pages - 1:
                document.add_page_break()
                
    document.save(str(output_path))


def convert_pdf_to_word(
    pdf_path: str | Path,
    output_path: str | Path | None = None,
    *,
    ocr_fallback: bool = True,
    ocr_lang: str = 'eng',
) -> Path:
    """
    Convert a system input PDF to a strictly formatted localized DOCX output.
    Actively routes extraction processes across hybrid parsing, standard conversions, 
    and pure optical OCR processing evaluation matrices.

    Args:
        pdf_path: Application-level input absolute context structure definition.
        output_path: Configured application-level output context routing element definition context wrapper.
        ocr_fallback: Flag toggling raw strict standard Tesseract OCR optical routing output bypass algorithms natively.
        ocr_lang: Tesseract string parameter linguistic configuration constraint string output.

    Returns:
        Path: Instantiated verified document mapping object reference variable definition context structure variable wrapper layout layout structure context map output definition.
    """
    source_path: Path = Path(pdf_path)
    validate_pdf_file(source_path)

    if output_path is None:
        output_path = source_path.with_suffix('.docx')

    target_path: Path = Path(output_path)
    if target_path.suffix.lower() != '.docx':
        target_path = target_path.with_suffix('.docx')

    target_path.parent.mkdir(parents=True, exist_ok=True)

    if ocr_fallback and not is_text_pdf(source_path):
        _create_docx_from_ocr(source_path, target_path, ocr_lang)
        return target_path

    try:
        _convert_hybrid(source_path, target_path, ocr_lang)
    except Exception as e:
        logger.error(f"Structured Premium Hybrid translation structural extraction constraint validation parsing error format exception mapping bypass fault engine bypassed fallback activated: {e}")
        try:
            _convert_pdf2docx(source_path, target_path)
            if not _docx_contains_text(target_path):
                raise ConversionError('Fallback validation map definition standard component standard standard evaluation returned format missing bounds fault elements output generation matrix error return matrix flag bypass string layout structure mapping validation elements strings dictionary zero.')
        except Exception as e2:
            logger.error(f"Secondary converter constraint fallback bounds bypass format bypass fault parsing string component exception error: {e2}")
            if not ocr_fallback:
                raise ConversionError(f"Complete system architectural translation bypass format error format error exception sequence block mapping block layout array structure format map structural bounds definitions bypassed constraints dictionary structure. Hybrid component bypass limit boundary strings structure bypassed exception matrix error format string structure formats formats dictionaries mapping matrix strings errors array definitions error component: {e2}") from e2
            
            try:
                _create_docx_from_ocr(source_path, target_path, ocr_lang)
            except Exception as e3:
                 raise ConversionError(f"Complete system conversion structure mapping array definition sequence structural error bypass framework sequence bounds mapping structure sequence formats dictionary component string array framework component boundaries exception block error limit arrays context arrays errors: OCR Error: {e3}") from e3

    return target_path
